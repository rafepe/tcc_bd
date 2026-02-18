from .models import *
from .tables import *
from contrapartida.models import *

import os
import zipfile
import unicodedata
from io import BytesIO
from itertools import groupby
from threading import Lock
from datetime import datetime, date
from decimal import Decimal, ROUND_HALF_UP

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

from django_tables2 import RequestConfig

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Sum
from django.http import HttpResponse, Http404, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse, reverse_lazy
from django.views.decorators.cache import never_cache
from django.views.generic import TemplateView
from django.views.generic.edit import DeleteView


def gerar_declaracoes(request):
    nome_projeto = request.GET.get('projeto')
    mes = request.GET.get('mes')
    ano = request.GET.get('ano')

    projeto_selecionado = None
    if nome_projeto:
        projeto_selecionado = projeto.objects.filter(nome=nome_projeto).first()

    contexto = {
        'projeto': projeto_selecionado,
        'mes': mes,
        'ano': ano,
    }

    return render(request, 'declaracao/gerar_declaracoes.html', contexto)

def declaracoes_menu(request):
    from datetime import date
    
    nome_projeto = request.GET.get('projeto')
    mes = request.GET.get('mes')
    ano = request.GET.get('ano')
    
    hoje = date.today()
    ano_atual = hoje.year
    semestre_atual = 1 if hoje.month <= 6 else 2

    if semestre_atual == 1:
        semestre = 2
        ano_default = ano_atual - 1
    else:
        semestre = 1
        ano_default = ano_atual

    ano = int(request.GET.get('ano', ano_default))
    semestre = int(request.GET.get('semestre', semestre))

    if semestre == 1:
        data_ini_semestre = date(ano, 1, 1)
        data_fim_semestre = date(ano, 6, 30)
    else:
        data_ini_semestre = date(ano, 7, 1)
        data_fim_semestre = date(ano, 12, 31)

    projetos = projeto.objects.filter(
        ativo=True,
        data_inicio__lte=data_fim_semestre,
        data_fim__gte=data_ini_semestre
    ).order_by('nome')

    projeto_obj = None
    if nome_projeto:
        projeto_obj = projeto.objects.filter(nome=nome_projeto).first()

    declaracoes = {}
    if projeto_obj and mes and ano:
        declaracoes = {
            'pesquisa': declaracao_contrapartida_pesquisa.objects.filter(id_projeto=projeto_obj.id, mes=mes, ano=ano).first(),
            'rh': declaracao_contrapartida_rh.objects.filter(id_projeto=projeto_obj.id, mes=mes, ano=ano).first(),
            'so': declaracao_contrapartida_so.objects.filter(id_projeto=projeto_obj.id, mes=mes, ano=ano).first(),
            'equipamento': declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).first(),
        }

    tipos = ["pesquisa","so","rh","equipamento"]    

    contexto = {
        'projeto': projeto_obj,
        'mes': mes,
        'ano': ano,
        'semestre': semestre,
        'declaracoes': declaracoes,
        'tipos': tipos,
        'projetos': projetos,
    }

    return render(request, 'declaracao/declaracoes_menu.html', contexto)

def gerar_declaracao_contrapartida_pesquisa(request, projeto_id, mes, ano):

    try:
        projeto_selecionado = projeto.objects.get(id=projeto_id)
    except projeto.DoesNotExist:
        messages.error(request, "projeto não encontrado.")
        return redirect('declaracoes_menu')

    declaracao_existente = declaracao_contrapartida_pesquisa.objects.filter(
        projeto=projeto_selecionado, mes=mes, ano=ano
    ).first()

    if declaracao_existente:
        messages.info(request, f"Já existe uma declaração para {projeto_selecionado.nome} - {mes}/{ano}.")
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

    declaracao = declaracao_contrapartida_pesquisa.objects.create(
        id_projeto=projeto_id,
        projeto=projeto_selecionado.nome,
        codigo=projeto_selecionado.peia,
        mes=mes,    	
        ano=ano,
    )

    registros = contrapartida_pesquisa.objects.filter(
        id_projeto=projeto_selecionado,
        id_salario__mes=mes,
        id_salario__ano=ano
    )

    if not registros.exists():
        messages.warning(request, "Nenhum dado encontrado para gerar a declaração.")
        declaracao.delete()
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')
    
    total=0

    for r in registros:
        pessoa_obj = r.id_salario.id_pessoa
        declaracao_contrapartida_pesquisa_item.objects.create(
            declaracao=declaracao,
            nome=pessoa_obj.nome,
            cpf=pessoa_obj.cpf,
            funcao=r.funcao,
            horas_alocadas=r.horas_alocadas,
            salario=r.id_salario.valor,
            valor_cp=r.valor_cp
        )

        total += r.valor_cp
    declaracao.total = total
    declaracao.save()

    messages.success(request, f"Declaração gerada para {projeto_selecionado.nome} - {mes}/{ano}.")
    url = reverse('declaracoes_menu')
    return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

class declaracao_contrapartida_pesquisa_view(TemplateView):
    template_name = 'declaracao/declaracao_contrapartida_pesquisa.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        id_declaracao = self.kwargs.get('id_declaracao')
        declaracao = get_object_or_404(declaracao_contrapartida_pesquisa, id=id_declaracao)

        itens = declaracao_contrapartida_pesquisa_item.objects.filter(declaracao=declaracao)

        tabela = declaracao_contrapartida_pesquisa_item_table(itens)
        RequestConfig(self.request).configure(tabela)

        context['declaracao'] = declaracao
        context['tabela'] = tabela

        return context

class declaracao_contrapartida_pesquisa_delete(DeleteView):

    def dispatch(self, request, *args, **kwargs):
        if request.user.has_perm("declaracao.declaracao_contrapartida_pesquisa_delete"):
            return super().dispatch(request, *args, **kwargs)
        else:
            return HttpResponse("Sem permissão para excluir declaracoes de contrapartida_pesquisa")

    model = declaracao_contrapartida_pesquisa
    fields = []
    template_name_suffix = '_delete'
    context_object_name = 'declaracao'
    success_url='/declaracao/menu/?projeto={projeto}&mes={mes}&ano={ano}'

    def get_success_url(self):
        return self.success_url.format(
            projeto=self.object.projeto,
            mes=self.object.mes,
            ano=self.object.ano
        )

def gerar_declaracao_contrapartida_so(request, projeto_id, mes, ano):

    try:
        projeto_selecionado = projeto.objects.get(id=projeto_id)
    except projeto.DoesNotExist:
        messages.error(request, "projeto não encontrado.")
        return redirect('declaracoes_menu')

    declaracao_existente = declaracao_contrapartida_so.objects.filter(
        projeto=projeto_selecionado, mes=mes, ano=ano
    ).first()

    if declaracao_existente:
        messages.info(request, f"Já existe uma declaração para {projeto_selecionado.nome} - {mes}/{ano}.")
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

    declaracao = declaracao_contrapartida_so.objects.create(
        id_projeto=projeto_id,
        projeto=projeto_selecionado.nome,
        codigo=projeto_selecionado.peia,
        mes=mes,    	
        ano=ano,
    )


    registros = contrapartida_so_projeto.objects.filter(
        id_projeto=projeto_selecionado,
        mes=mes,
        ano=ano
    )

    if not registros.exists():
        messages.warning(request, "Nenhum dado encontrado para gerar a declaração.")
        declaracao.delete()
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')
    
    total=0

    for r in registros:
        total += r.valor
    declaracao.total = total
    declaracao.save()

    messages.success(request, f"Declaração gerada para {projeto_selecionado.nome} - {mes}/{ano}.")
    url = reverse('declaracoes_menu')
    return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

class declaracao_contrapartida_so_view(TemplateView):
    template_name = 'declaracao/declaracao_contrapartida_so.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        id_declaracao = self.kwargs.get('id_declaracao')
        declaracao = get_object_or_404(declaracao_contrapartida_so, id=id_declaracao)
     
        context['declaracao'] = declaracao
        return context

class declaracao_contrapartida_so_delete(DeleteView):

    def dispatch(self, request, *args, **kwargs):
        if request.user.has_perm("declaracao.declaracao_contrapartida_so_delete"):
            return super().dispatch(request, *args, **kwargs)
        else:
            return HttpResponse("Sem permissão para excluir declaracoes de contrapartida_so")

    model = declaracao_contrapartida_so
    fields = []
    template_name_suffix = '_delete'
    context_object_name = 'declaracao'
    success_url='/declaracao/menu/?projeto={projeto}&mes={mes}&ano={ano}'

    def get_success_url(self):
        return self.success_url.format(
            projeto=self.object.projeto,
            mes=self.object.mes,
            ano=self.object.ano
        )

def gerar_declaracao_contrapartida_rh(request, projeto_id, mes, ano):

    try:
        projeto_selecionado = projeto.objects.get(id=projeto_id)
    except projeto.DoesNotExist:
        messages.error(request, "projeto não encontrado.")
        return redirect('declaracoes_menu')

    declaracao_existente = declaracao_contrapartida_rh.objects.filter(
        projeto=projeto_selecionado.nome, mes=mes, ano=ano
    ).first()

    if declaracao_existente:
        messages.info(request, f"Já existe uma declaração para {projeto_selecionado.nome} - {mes}/{ano}.")
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

    declaracao = declaracao_contrapartida_rh.objects.create(
        id_projeto=projeto_id,
        projeto=projeto_selecionado.nome,
        codigo=projeto_selecionado.peia,
        mes=mes,    	
        ano=ano,
    )

    registros = contrapartida_rh.objects.filter(
        id_projeto=projeto_selecionado,
        id_salario__mes=mes,
        id_salario__ano=ano
    )

    if not registros.exists():
        messages.warning(request, "Nenhum dado encontrado para gerar a declaração.")
        declaracao.delete()
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')
    total=0

    for r in registros:
        pessoa_obj = r.id_salario.id_pessoa
        declaracao_contrapartida_rh_item.objects.create(
            declaracao=declaracao,
            nome=pessoa_obj.nome,
            cpf=pessoa_obj.cpf,
            funcao=r.funcao,
            horas_alocadas=r.horas_alocadas,
            salario=r.id_salario.valor,
            salario_mes=r.id_salario.mes,
            salario_ano=r.id_salario.ano,
            valor_cp=r.valor_cp,
            valor_hora=r.id_salario.valor_hora
        )

        total += r.valor_cp
    declaracao.total = total
    declaracao.save()

    messages.success(request, f"Declaração gerada para {projeto_selecionado.nome} - {mes}/{ano}.")
    url = reverse('declaracoes_menu')
    return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

class declaracao_contrapartida_rh_view(TemplateView):
    template_name = 'declaracao/declaracao_contrapartida_rh.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        id_declaracao = self.kwargs.get('id_declaracao')
        declaracao = get_object_or_404(declaracao_contrapartida_rh, id=id_declaracao)

        itens = declaracao_contrapartida_rh_item.objects.filter(declaracao=declaracao)

        tabela = declaracao_contrapartida_rh_item_table(itens)
        RequestConfig(self.request).configure(tabela)

        context['declaracao'] = declaracao
        context['tabela'] = tabela

        return context

class declaracao_contrapartida_rh_delete(DeleteView):

    def dispatch(self, request, *args, **kwargs):
        if request.user.has_perm("declaracao.declaracao_contrapartida_rh_delete"):
            return super().dispatch(request, *args, **kwargs)
        else:
            return HttpResponse("Sem permissão para excluir declaracoes de contrapartida_rh")

    model = declaracao_contrapartida_rh
    fields = []
    template_name_suffix = '_delete'
    context_object_name = 'declaracao'
    success_url='/declaracao/menu/?projeto={projeto}&mes={mes}&ano={ano}'

    def get_success_url(self):
        return self.success_url.format(
            projeto=self.object.projeto,
            mes=self.object.mes,
            ano=self.object.ano
        )

def gerar_declaracao_contrapartida_equipamento(request, projeto_id, mes, ano):
    try:
        projeto_selecionado = projeto.objects.get(id=projeto_id)
    except projeto.DoesNotExist:
        messages.error(request, "projeto não encontrado.")
        return redirect('declaracoes_menu')

    declaracao_existente = declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).first()
    if declaracao_existente:
        messages.info(request, f"Já existe uma declaração para {mes}/{ano}.")
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

    declaracao = declaracao_contrapartida_equipamento.objects.create(mes=mes, ano=ano)

    registros = (
        contrapartida_equipamento.objects
        .filter(mes=mes, ano=ano)
        .select_related("id_projeto", "id_equipamento")
        .order_by("id_equipamento__nome", "id_projeto__id")
    )

    if not registros.exists():
        messages.warning(request, "Nenhum dado encontrado para gerar a declaração.")
        declaracao.delete()
        url = reverse('declaracoes_menu')
        return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')

    for (equip_nome, proj_id), grupo in groupby(
        registros,
        key=lambda r: (r.id_equipamento.nome, r.id_projeto_id)
    ):
        grupo = list(grupo)

        horas_total = sum(Decimal(g.horas_alocadas or 0) for g in grupo)
        valor_total = sum(Decimal(g.valor_cp or 0) for g in grupo)

        # junta descrições únicas (se quiser)
        descricoes = []
        for g in grupo:
            d = (g.descricao or "").strip()
            if d and d not in descricoes:
                descricoes.append(d)
        descricao_final = "; ".join(descricoes)

        p0 = grupo[0].id_projeto

        declaracao_contrapartida_equipamento_item.objects.create(
            declaracao=declaracao,
            codigo=p0.peia,
            projeto=p0.nome,
            equipamento=equip_nome,
            descricao=descricao_final,
            horas_alocadas=horas_total,   # se seu campo for IntegerField, use int(horas_total)
            valor_cp=valor_total
        )

    messages.success(request, f"Declaração gerada para {mes}/{ano}.")
    url = reverse('declaracoes_menu')
    return redirect(f'{url}?projeto={projeto_selecionado.nome}&mes={mes}&ano={ano}')


class declaracao_contrapartida_equipamento_view(TemplateView):
    template_name = 'declaracao/declaracao_contrapartida_equipamento.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        id = self.kwargs.get('id_declaracao') 
        declaracao = get_object_or_404(declaracao_contrapartida_equipamento, id=id)

        itens = declaracao_contrapartida_equipamento_item.objects.filter(declaracao=declaracao)

        tabela = declaracao_contrapartida_equipamento_item_table(itens)
        RequestConfig(self.request).configure(tabela)

        context['declaracao'] = declaracao
        context['tabela'] = tabela

        return context

class declaracao_contrapartida_equipamento_delete(DeleteView):

    def dispatch(self, request, *args, **kwargs):
        if request.user.has_perm("declaracao.declaracao_contrapartida_equipamento_delete"):
            return super().dispatch(request, *args, **kwargs)
        else:
            return HttpResponse("Sem permissão para excluir declaracoes de contrapartida_equipamento")

    model = declaracao_contrapartida_equipamento
    fields = []
    template_name_suffix = '_delete'
    context_object_name = 'declaracao'

    def get_success_url(self):
        return reverse_lazy('declaracoes_menu')

def gerar_docx_rh(request, declaracao_id):
    declaracao_itens = declaracao_contrapartida_rh_item.objects.filter(declaracao_id=declaracao_id)
    if not declaracao_itens.exists():
        return HttpResponse("Nenhuma declaração encontrada para esse mês e ano.")

    ano = declaracao_itens.first().salario_ano
    mes = declaracao_itens.first().salario_mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()

    
    if declaracao_itens.exists():
        projeto_nome = declaracao_itens.first().declaracao.projeto
        total_valor_cp = declaracao_itens.aggregate(total=Sum('valor_cp'))['total'] or 0

    caminho_docx = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    doc = Document(caminho_docx)

    for p in doc.paragraphs:
        p.text = (
            p.text
            .replace('{{mes_selecionado}}', mes_nome)
            .replace('{{ano_selecionado}}', str(ano))
            .replace('{{nome_projeto}}', projeto_nome)
            .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
        )
    itens = list(declaracao_itens) if declaracao_itens.exists() else []
   
    for i, paragraph in enumerate(doc.paragraphs):
        if '{{tabela_itens}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')

            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nome do Bolsista'
            hdr_cells[1].text = 'CPF'
            hdr_cells[2].text = 'Função'
            hdr_cells[3].text = 'Horas alocadas'
            hdr_cells[4].text = 'Salário'
            hdr_cells[5].text = 'Valor CP'
            hdr_cells[6].text = 'Valor Hora'

            for item in itens:
                row_cells = table.add_row().cells
                row_cells[0].text = item.nome
                row_cells[1].text = item.cpf
                row_cells[2].text = item.funcao
                row_cells[3].text = str(item.horas_alocadas)
                row_cells[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row_cells[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row_cells[6].text = f"R$ {item.valor_hora:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            break  
    p = doc.add_paragraph('(*) Valor das horas é o produto da multiplicação entre o nº de horas e o quociente da divisão do valor do salário por 160.')
    p = doc.add_paragraph('(**) Mês da competência do contracheque.')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')
    p = doc.add_paragraph('')

    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    row = table.rows[0]
    row.cells[0].text = "Anderson Soares"
    row.cells[1].text = "Telma Woerle de Lima Soares"

    row = table.rows[1]
    row.cells[0].text = "Coordenador do Projeto"
    row.cells[1].text = "Diretora da Unidade Embrapii - CEIA/UFG"

    for r in table.rows:
        for cell in r.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'nil')
            tcPr.append(element)



    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="RH_{projeto_nome}_{mes}_{ano}.docx"'
    doc.save(response)
    return response

try:
    from num2words import num2words
except ImportError:
    raise ImportError("Instale a dependência: pip install num2words")

def valor_por_extenso(valor):
    """
    Recebe float/Decimal e devolve, em pt-BR, ex.: 'doze reais e trinta e quatro centavos'
    """
    v = Decimal(valor).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    return num2words(v, to='currency', lang='pt_BR')

def gerar_docx_so(request, declaracao_id):
    declaracao = get_object_or_404(declaracao_contrapartida_so, id=declaracao_id)

    ano = declaracao.ano
    mes = declaracao.mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    codigo_peia = declaracao.codigo or ''

    total_valor_cp = declaracao.total or 0.0

    caminho_base_preferido = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_so.docx')
    caminho_fallback = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    caminho_docx = caminho_base_preferido if os.path.exists(caminho_base_preferido) else caminho_fallback
    doc = Document(caminho_docx)

    def br_currency(v):
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def apply_replacements_to_text(text):
        return (text
                .replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{codigo_peia}}', codigo_peia)
                .replace('{{valor_total}}', br_currency(total_valor_cp))
                .replace('{{valor_extenso}}', valor_por_extenso(total_valor_cp)))

    for p in doc.paragraphs:
        if p.text:
            p.text = apply_replacements_to_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = apply_replacements_to_text(p.text)

    marcador_encontrado = False
    
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="SO_{projeto_nome}_{mes}_{ano}.docx"'
    doc.save(response)
    return response

def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def _set_cell_borders(cell, color="000000", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), size)
        tag.set(qn('w:space'), "0")
        tag.set(qn('w:color'), color)
        tc_borders.append(tag)

def _fmt_moeda_br_decimal(v: Decimal | float | int | None) -> str:
    if v is None:
        v = Decimal("0")
    v = Decimal(v)
    s = f"{v:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {s}"

def _nota_equipamento_por_nome(nome_equip: str) -> str | None:
    eq = equipamento.objects.filter(nome__iexact=nome_equip).first()
    if not eq:
        return None
    aquis = eq.valor_aquisicao or Decimal("0")
    nos = eq.quantidade_nos or 1
    custo_unid = Decimal(eq.valor_hora or 0)
    unidade = "GPU" if eq.nome.upper().startswith("DGX") else "nó"
    return (
        f"valor de compra original de {_fmt_moeda_br_decimal(aquis)} "
        f"e um custo por {unidade} de {_fmt_moeda_br_decimal(custo_unid)} "
        f"onde o equipamento conta com {nos} {unidade + ('s' if nos != 1 else '')}"
    )

def gerar_docx_equipamento(request, id):
    declaracao = get_object_or_404(declaracao_contrapartida_equipamento, id=id)
    itens = declaracao.itens.all().order_by("equipamento", "projeto", "codigo")
    if not itens.exists():
        raise Http404("Não há itens para esta declaração.")
    base_path = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_equipamento.docx")
    doc = Document(base_path)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(4.63)  # distância da borda ao cabeçalho


    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)

    meses = ["", "JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO",
             "JULHO","AGOSTO","SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"]
    
    p = doc.add_paragraph("DECLARAÇÃO DE USO DE EQUIPAMENTOS")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"MÊS DE COMPETÊNCIA: {meses[declaracao.mes]} DE {declaracao.ano}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")

    col_widths = [Cm(3.5), Cm(5.5), Cm(2.5), Cm(8.0), Cm(3.5)]

    for equip_nome, grupo in groupby(itens, key=lambda i: i.equipamento):
        grupo = list(grupo)
        nota = _nota_equipamento_por_nome(equip_nome)

        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"
        table.autofit = True

        faixa = table.rows[0].cells[0].merge(
            table.rows[0].cells[1]
        ).merge(
            table.rows[0].cells[2]
        ).merge(
            table.rows[0].cells[3]
        ).merge(
            table.rows[0].cells[4]
        )
        _set_cell_shading(faixa, "1F4E79")
        _set_cell_borders(faixa, color="1F4E79")
        par = faixa.paragraphs[0]
        run = par.add_run(f"EQUIPAMENTO: {equip_nome}")
        if nota:
            run.add_text(" *")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

        headers = ["CÓDIGO", "TÍTULO DO PROJETO", "HORAS NO MÊS",
                   "DESCRIÇÃO DAS ATIVIDADES", "VALOR DA CONTRAPARTIDA"]
        for j, text in enumerate(headers):
            c = table.rows[1].cells[j]
            c.width = col_widths[j]
            _set_cell_shading(c, "D9D9D9")
            _set_cell_borders(c)
            par = c.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(text)
            run.bold = True

        total_equip = Decimal("0")
        for item in grupo:
            row = table.add_row()
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
            row.cells[0].text = item.codigo or ""
            row.cells[1].text = item.projeto or ""
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].text = str(item.horas_alocadas or 0)
            row.cells[3].text = item.descricao or ""
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[4].text = _fmt_moeda_br_decimal(item.valor_cp or 0)
            for c in row.cells: _set_cell_borders(c)
            total_equip += Decimal(item.valor_cp or 0)

        total_row = table.add_row()
        merged = total_row.cells[0].merge(total_row.cells[1]).merge(
                 total_row.cells[2]).merge(total_row.cells[3])
        _set_cell_shading(merged, "EFEFEF")
        _set_cell_borders(merged)
        merged.paragraphs[0].add_run("TOTAL DA CONTRAPARTIDA").bold = True
        val_cell = total_row.cells[4]
        _set_cell_shading(val_cell, "EFEFEF")
        _set_cell_borders(val_cell)
        par = val_cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        par.add_run(_fmt_moeda_br_decimal(total_equip)).bold = True

        if nota:
            doc.add_paragraph(f"(*) {nota}").runs[0].italic = True

        doc.add_paragraph("")

    total_geral = sum(Decimal(i.valor_cp or 0) for i in itens)
    par = doc.add_paragraph("TOTAL GERAL DA CONTRAPARTIDA: ")
    par.runs[0].bold = True
    par.add_run(_fmt_moeda_br_decimal(total_geral)).bold = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = (
        f'attachment; filename="{declaracao.ano}-{declaracao.mes:02d}-Equipamentos.docx"'
    )
    return resp

def gerar_docx_pesquisa(request, declaracao_id):
    # Busca a declaração e itens
    declaracao = get_object_or_404(declaracao_contrapartida_pesquisa, id=declaracao_id)
    itens_qs = declaracao.itens.all()

    if not itens_qs.exists():
        return HttpResponse("Nenhum item encontrado para esta declaração de Pesquisa.")

    # Metadados (mês/ano/projeto)
    ano = declaracao.ano
    mes = declaracao.mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto

    total_valor_cp = itens_qs.aggregate(total=Sum('valor_cp'))['total'] or 0

    # Abre o DOCX base (use um arquivo próprio p/ Pesquisa; cai no base_rh.docx se não existir)
    caminho_base_preferido = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_pesquisa.docx')
    caminho_fallback = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    caminho_docx = caminho_base_preferido if os.path.exists(caminho_base_preferido) else caminho_fallback

    doc = Document(caminho_docx)

    # Substituição simples em parágrafos
    for p in doc.paragraphs:
        if p.text:
            p.text = (
                p.text
                .replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
            )

    # Também substitui placeholders em células de tabelas, se existirem
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = (
                            p.text
                            .replace('{{mes_selecionado}}', mes_nome)
                            .replace('{{ano_selecionado}}', str(ano))
                            .replace('{{nome_projeto}}', projeto_nome)
                            .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
                        )

    # Insere a tabela de itens no marcador {{tabela_itens}}
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def inserir_tabela_itens(ap_depois):
        # Cria tabela 6 colunas: Nome | CPF | Função | Horas | Salário | Valor CP
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Nome do Pesquisador'
        hdr[1].text = 'CPF'
        hdr[2].text = 'Função'
        hdr[3].text = 'Horas alocadas'
        hdr[4].text = 'Salário'
        hdr[5].text = 'Valor CP'

        for item in itens_qs:
            row = table.add_row().cells
            row[0].text = item.nome
            row[1].text = item.cpf
            row[2].text = item.funcao
            row[3].text = str(item.horas_alocadas)
            row[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        ap_depois._element.addnext(table._tbl)

    marcador_encontrado = False
    for paragraph in doc.paragraphs:
        if '{{tabela_itens}}' in paragraph.text:
            marcador_encontrado = True
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')
            inserir_tabela_itens(paragraph)
            break

    if not marcador_encontrado:
        p_anchor = doc.add_paragraph()
        inserir_tabela_itens(p_anchor)

    doc.add_paragraph('(*) Valor das horas é o produto da multiplicação entre o nº de horas e o quociente da divisão do valor do salário por 160.')
    doc.add_paragraph('(**) Mês da competência do contracheque.')

    # Tabela de assinaturas
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    table_ass = doc.add_table(rows=2, cols=2)
    table_ass.style = 'Table Grid'

    table_ass.rows[0].cells[0].text = "Anderson Soares"
    table_ass.rows[0].cells[1].text = "Telma Woerle de Lima Soares"
    table_ass.rows[1].cells[0].text = "Coordenador do Projeto"
    table_ass.rows[1].cells[1].text = "Diretora da Unidade Embrapii - CEIA/UFG"

    for r in table_ass.rows:
        for cell in r.cells:
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Remove bordas da tabela de assinatura
    tbl = table_ass._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'nil')
            tcPr.append(element)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="PESQUISA_{projeto_nome}_{mes}_{ano}.docx"'
    doc.save(response)
    return response

def central_declaracoes(request):
    projeto_id = request.GET.get('projeto_id')
    ano = request.GET.get('ano')
    semestre = request.GET.get('semestre')

    if not (projeto_id and ano and semestre):
        meses_data = []
    else:
        meses_semestre = [1, 2, 3, 4, 5, 6] if semestre == '1' else [7, 8, 9, 10, 11, 12]

        meses_data = []
        for mes in meses_semestre:
            registros_mes = declaracao_contrapartida_so.objects.filter(
                projeto_id=projeto_id,
                data__year=ano,
                data__month=mes,
            )

            meses_data.append({
                'mes': mes,
                'contrapartidas': registros_mes, 
            })

    context = {
        'meses_data': meses_data,
        'ano': ano,
        'projeto_id': projeto_id,
        'semestre': semestre,
    }
    return render(request, 'declaracao/central.html', context)

def visualizar_declaracao(request, declaracao_id):
    """
    View para visualizar uma declaração específica
    """
    declaracao = get_object_or_404(declaracao_contrapartida_rh, id=declaracao_id)
    itens = declaracao_contrapartida_rh_item.objects.filter(declaracao=declaracao).order_by('nome')
    
    context = {
        'declaracao': declaracao,
        'itens': itens,
    }
    
    return render(request, 'declaracao/visualizar.html', context)

def download_declaracao(request, declaracao_id):
    """
    View para download da declaração em PDF (implementação básica)
    """
    declaracao = get_object_or_404(declaracao_contrapartida_rh, id=declaracao_id)
    itens = declaracao_contrapartida_rh_item.objects.filter(declaracao=declaracao).order_by('nome')
    
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from django.http import HttpResponse
        import io
        
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        # Cabeçalho
        p.drawString(100, 750, f"Declaração de Contrapartida RH")
        p.drawString(100, 730, f"Projeto: {declaracao.projeto}")
        p.drawString(100, 710, f"Período: {declaracao.mes}/{declaracao.ano}")
        p.drawString(100, 690, f"Total: R$ {declaracao.total:.2f}")
        
        # Itens
        y_position = 650
        for item in itens:
            p.drawString(100, y_position, f"{item.nome} - {item.funcao} - R$ {item.valor_cp:.2f}")
            y_position -= 20
            if y_position < 100:
                p.showPage()
                y_position = 750
        
        p.showPage()
        p.save()
        
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="declaracao_{declaracao.id}.pdf"'
        return response
        
    except ImportError:
        messages.warning(request, "Biblioteca de PDF não disponível. Mostrando versão HTML.")
        return render(request, 'declaracao/pdf.html', {
            'declaracao': declaracao,
            'itens': itens,
        })

def editar_declaracao(request, declaracao_id):
    """
    View para editar uma declaração
    """
    declaracao = get_object_or_404(declaracao_contrapartida_rh, id=declaracao_id)
    itens = declaracao_contrapartida_rh_item.objects.filter(declaracao=declaracao).order_by('nome')
    
    if request.method == 'POST':
        for item in itens:
            novo_valor = request.POST.get(f'valor_cp_{item.id}')
            if novo_valor:
                try:
                    item.valor_cp = float(novo_valor)
                    item.save()
                except ValueError:
                    pass
        
        declaracao.total = sum(item.valor_cp for item in itens)
        declaracao.save()
        
        messages.success(request, "Declaração atualizada com sucesso!")
        return redirect('visualizar_declaracao', declaracao_id=declaracao.id)
    
    context = {
        'declaracao': declaracao,
        'itens': itens,
    }
    
    return render(request, 'declaracao/editar.html', context)

def redirect_to_central(projeto_id, ano, semestre, projeto_nome=None, mes=None):
    """
    Função auxiliar para redirecionamento para a central
    """
    url = reverse('central_declaracoes')
    params = f'?projeto_id={projeto_id}&ano={ano}&semestre={semestre}'
    
    if projeto_nome:
        params += f'&projeto={projeto_nome}'
    if mes:
        params += f'&mes={mes}'
        
    return redirect(f'{url}{params}')

def get_nome_mes(numero_mes):
    
    """
    Função auxiliar para converter número do mês em nome
    """
    meses = {
        1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
        5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
        9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
    }
    return meses.get(numero_mes, f'Mês {numero_mes}')

@login_required
def ajax_dados_projeto(request):
    """
    View AJAX para carregar dados de um projeto específico por semestre
    """
    projeto_id = request.GET.get('projeto_id')
    semestre = int(request.GET.get('semestre', 1))
    ano = int(request.GET.get('ano'))
    
    if not projeto_id:
        return JsonResponse({'error': 'Projeto não especificado'}, status=400)
    
    try:
        projeto_obj = projeto.objects.get(id=projeto_id)
    except projeto.DoesNotExist:
        return JsonResponse({'error': 'Projeto não encontrado'}, status=404)
    
    meses_semestre = [1, 2, 3, 4, 5, 6] if semestre == 1 else [7, 8, 9, 10, 11, 12]
    
    dados_meses = {}
    
    for mes in meses_semestre:
        dados_meses[mes] = {
            'rh': verificar_declaracao_rh(projeto_obj.id, mes, ano),
            'pesquisa': verificar_declaracao_pesquisa(projeto_obj.id, mes, ano),
            'so': verificar_declaracao_so(projeto_obj.id, mes, ano),
            'equipamentos': verificar_declaracao_equipamentos(mes, ano),
            }
    
    resposta = {
        'projeto': {
            'id': projeto_obj.id,
            'nome': projeto_obj.nome,
            'codigo': projeto_obj.peia,
        },
        'semestre': semestre,
        'ano': ano,
        'meses': dados_meses
    }
    
    return JsonResponse(resposta)

def verificar_declaracao_rh(projeto_obj, mes, ano):
    """
    Verifica status da declaração de contrapartida RH para um mês específico
    """
    declaracao_existente = declaracao_contrapartida_rh.objects.filter(
        id_projeto=projeto_obj,
        mes=mes,
        ano=ano
    ).first()
    
    if declaracao_existente:
        return {
            'existe': True,
            'pode_gerar': False,
            'valor': float(declaracao_existente.total or 0),
            'declaracao_id': declaracao_existente.id,
            'data_criacao': declaracao_existente.data_geracao.strftime('%d/%m/%Y %H:%M') if declaracao_existente.data_geracao else None
        }
    
    registros_contrapartida = contrapartida_rh.objects.filter(
        id_projeto=projeto_obj,
        id_salario__mes=mes,
        id_salario__ano=ano
    )
    
    pode_gerar = registros_contrapartida.exists()
    valor_potencial = sum(r.valor_cp for r in registros_contrapartida) if pode_gerar else 0
    
    return {
        'existe': False,
        'pode_gerar': pode_gerar,
        'valor': 0,
        'valor_potencial': float(valor_potencial),
        'quantidade_registros': registros_contrapartida.count()
    }

def verificar_declaracao_pesquisa(projeto_obj, mes, ano):
    """
    Verifica status da declaração de contrapartida Pesquisa para um mês específico
    """
    declaracao_existente = declaracao_contrapartida_pesquisa.objects.filter(
        id_projeto=projeto_obj,
        mes=mes,
        ano=ano
    ).first()
    if declaracao_existente:
        return {
            'existe': True,
            'pode_gerar': False,
            'valor': float(declaracao_existente.total or 0),
            'declaracao_id': declaracao_existente.id,
            'data_criacao': declaracao_existente.data_geracao.strftime('%d/%m/%Y %H:%M') if declaracao_existente.data_geracao else None
        }
    
    registros_contrapartida = contrapartida_pesquisa.objects.filter(
        id_projeto=projeto_obj,
        id_salario__mes=mes,
        id_salario__ano=ano
    )
    
    pode_gerar = registros_contrapartida.exists()
    valor_potencial = sum(r.valor_cp for r in registros_contrapartida) if pode_gerar else 0
    
    return {
        'existe': False,
        'pode_gerar': pode_gerar,
        'valor': 0,
        'valor_potencial': float(valor_potencial),
        'quantidade_registros': registros_contrapartida.count()
    }

def verificar_declaracao_so(projeto_obj, mes, ano):
    """
    Verifica status da declaração de contrapartida SO para um mês específico
    """
    declaracao_existente = declaracao_contrapartida_so.objects.filter(
        id_projeto=projeto_obj,
        mes=mes,
        ano=ano
    ).first()
    
    if declaracao_existente:
        return {
            'existe': True,
            'pode_gerar': False,
            'valor': float(declaracao_existente.total or 0),
            'declaracao_id': declaracao_existente.id,
            'data_criacao': declaracao_existente.data_geracao.strftime('%d/%m/%Y %H:%M') if declaracao_existente.data_geracao else None
        }
    
    registros_contrapartida = contrapartida_so_projeto.objects.filter(
        id_projeto=projeto_obj,
        mes=mes,
        ano=ano
    )
    
    pode_gerar = registros_contrapartida.exists()
    valor_potencial = sum(r.valor for r in registros_contrapartida) if pode_gerar else 0
    
    return {
        'existe': False,
        'pode_gerar': pode_gerar,
        'valor': 0,
        'valor_potencial': float(valor_potencial),
        'quantidade_registros': registros_contrapartida.count()
    }

def verificar_declaracao_equipamentos(mes, ano):
    """
    Verifica status da declaração de contrapartida Equipamentos para um mês específico
    """
    declaracao_existente = declaracao_contrapartida_equipamento.objects.filter(
        mes=mes,
        ano=ano
    ).first()
    
    if declaracao_existente:
        itens_projeto = declaracao_contrapartida_equipamento_item.objects.filter(
            declaracao=declaracao_existente,
        )
        valor_total = sum(item.valor_cp for item in itens_projeto) if itens_projeto.exists() else 0
        
        return {
            'existe': itens_projeto.exists(),
            'pode_gerar': False,
            'valor': float(valor_total),
            'declaracao_id': declaracao_existente.id,
            'data_criacao': declaracao_existente.data_geracao.strftime('%d/%m/%Y %H:%M') if declaracao_existente.data_geracao else None
        }
    
    registros_contrapartida = contrapartida_equipamento.objects.filter(
        mes=mes,
        ano=ano
    )
    
    pode_gerar = registros_contrapartida.exists()
    valor_potencial = sum(r.valor_cp for r in registros_contrapartida) if pode_gerar else 0
    
    return {
        'existe': False,
        'pode_gerar': pode_gerar,
        'valor': 0,
        'valor_potencial': float(valor_potencial),
        'quantidade_registros': registros_contrapartida.count()
    }

@login_required
def download_declaracao_mes(request):
    """
    View para download de todas as declarações de um mês específico
    """
    projeto_id = request.GET.get('projeto_id')
    mes = int(request.GET.get('mes'))
    ano = int(request.GET.get('ano'))
    
    try:
        projeto_obj = projeto.objects.get(id=projeto_id)
    except projeto.DoesNotExist:
        return JsonResponse({'error': 'Projeto não encontrado'}, status=404)
    
    declaracoes = {
        'rh': declaracao_contrapartida_rh.objects.filter(
            id_projeto=projeto_id, mes=mes, ano=ano
        ).first(),
        'custeio': None,
        'equipamentos': None,
    }
    
    if declaracoes['rh']:
        return redirect('download_declaracao', declaracao_id=declaracoes['rh'].id)
    
    return JsonResponse({'error': 'Nenhuma declaração encontrada para este mês'}, status=404)

def gerar_docx_rh_novo(declaracao_id):
    """
    Gera DOCX de RH e salva no disco segundo o padrão:
    media/declaracoes/ANO-SEMESTRE/PEIA-NOME-rh-ANO-MES.docx
    """
    import os
    from datetime import datetime
    from docx import Document
    from django.conf import settings
    from django.http import Http404
    from django.db.models import Sum

    decl_itens = declaracao_contrapartida_rh_item.objects.filter(declaracao_id=declaracao_id)
    if not decl_itens.exists():
        raise Http404("Nenhuma declaração RH encontrada.")

    item0 = decl_itens.first()
    ano = item0.salario_ano
    mes = item0.salario_mes
    semestre = 1 if mes <= 6 else 2

    projeto_nome = item0.declaracao.projeto
    projeto_nome_slug = projeto_nome.replace(" ", "_")
    peia = item0.declaracao.codigo or "SEMPEIA"

    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()

    total_valor_cp = decl_itens.aggregate(total=Sum('valor_cp'))['total'] or 0

    # Template base existente
    path_template = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    doc = Document(path_template)

    # ===== REPLACE placeholders =====
    for p in doc.paragraphs:
        p.text = (
            p.text
            .replace('{{mes_selecionado}}', mes_nome)
            .replace('{{ano_selecionado}}', str(ano))
            .replace('{{nome_projeto}}', projeto_nome)
            .replace('{{valor_total}}',
                     f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
        )

    # ===== INSERT TABLE =====
    for paragraph in doc.paragraphs:
        if '{{tabela_itens}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            hdr = table.rows[0].cells
            hdr[0].text = 'Nome'
            hdr[1].text = 'CPF'
            hdr[2].text = 'Função'
            hdr[3].text = 'Horas'
            hdr[4].text = 'Salário'
            hdr[5].text = 'Valor CP'
            hdr[6].text = 'Valor Hora'

            for item in decl_itens:
                row = table.add_row().cells
                row[0].text = item.nome
                row[1].text = item.cpf
                row[2].text = item.funcao
                row[3].text = str(item.horas_alocadas)
                row[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row[6].text = f"R$ {item.valor_hora:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            break

    # ===== DESTINO DO ARQUIVO =====
    pasta_destino = os.path.join(
        settings.MEDIA_ROOT,
        "declaracoes",
        f"{ano}-{semestre}"
    )
    os.makedirs(pasta_destino, exist_ok=True)

    filename = f"{peia}-{projeto_nome_slug}-rh-{ano}-{mes:02d}.docx"
    caminho_final = os.path.join(pasta_destino, filename)

    doc.save(caminho_final)

    return caminho_final

def gerar_docx_pesquisa_novo(declaracao_id):
    """
    Salva:
    media/declaracoes/ANO-SEMESTRE/PEIA-NOME-pesquisa-AAAA-MM.docx
    """
    import os
    from datetime import datetime
    from django.conf import settings
    from django.http import Http404
    from django.db.models import Sum
    from docx import Document

    decl = get_object_or_404(declaracao_contrapartida_pesquisa, id=declaracao_id)
    itens = decl.itens.all()

    if not itens.exists():
        raise Http404("Nenhum item para esta declaração de Pesquisa.")

    ano = decl.ano
    mes = decl.mes
    semestre = 1 if mes <= 6 else 2

    projeto_nome = decl.projeto
    projeto_slug = projeto_nome.replace(" ", "_")
    peia = decl.codigo or "SEMPEIA"

    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()

    total_cp = itens.aggregate(total=Sum("valor_cp"))["total"] or 0

    # arquivo base
    path_template = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_pesquisa.docx")
    if not os.path.exists(path_template):
        path_template = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_rh.docx")

    doc = Document(path_template)

    # substituição simples
    def replace_all(text):
        return (
            text.replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{valor_total}}',
                         f"R$ {total_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
        )

    for p in doc.paragraphs:
        if p.text:
            p.text = replace_all(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = replace_all(p.text)

    # inserir tabela {{tabela_itens}}
    for paragraph in doc.paragraphs:
        if "{{tabela_itens}}" in paragraph.text:
            paragraph.text = ""
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            hdr[0].text = "Nome"
            hdr[1].text = "CPF"
            hdr[2].text = "Função"
            hdr[3].text = "Horas"
            hdr[4].text = "Salário"
            hdr[5].text = "Valor CP"

            for item in itens:
                row = table.add_row().cells
                row[0].text = item.nome
                row[1].text = item.cpf
                row[2].text = item.funcao
                row[3].text = str(item.horas_alocadas)
                row[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            break

    # destino final
    pasta = os.path.join(settings.MEDIA_ROOT, "declaracoes", f"{ano}-{semestre}")
    os.makedirs(pasta, exist_ok=True)

    filename = f"{peia}-{projeto_slug}-pesquisa-{ano}-{mes:02d}.docx"
    fullpath = os.path.join(pasta, filename)

    doc.save(fullpath)
    return fullpath

def gerar_docx_so_novo(declaracao_id):
    """
    Salva:
    media/declaracoes/ANO-SEMESTRE/PEIA-NOME-so-AAAA-MM.docx
    """
    import os
    from datetime import datetime
    from django.conf import settings
    from django.http import Http404
    from docx import Document

    decl = get_object_or_404(declaracao_contrapartida_so, id=declaracao_id)

    ano = decl.ano
    mes = decl.mes
    semestre = 1 if mes <= 6 else 2

    projeto_nome = decl.projeto
    projeto_slug = projeto_nome.replace(" ", "_")
    peia = decl.codigo or "SEMPEIA"

    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()

    total_cp = decl.total or 0.0

    # template base
    path_base = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_so.docx")
    if not os.path.exists(path_base):
        path_base = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_rh.docx")

    doc = Document(path_base)

    # função auxiliar
    def br_currency(v):
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def replace_all(text):
        return (text
                .replace("{{mes_selecionado}}", mes_nome)
                .replace("{{ano_selecionado}}", str(ano))
                .replace("{{nome_projeto}}", projeto_nome)
                .replace("{{codigo_peia}}", peia)
                .replace("{{valor_total}}", br_currency(total_cp))
                )

    # aplica
    for p in doc.paragraphs:
        if p.text:
            p.text = replace_all(p.text)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = replace_all(p.text)

    # destino
    pasta = os.path.join(settings.MEDIA_ROOT, "declaracoes", f"{ano}-{semestre}")
    os.makedirs(pasta, exist_ok=True)

    filename = f"{peia}-{projeto_slug}-so-{ano}-{mes:02d}.docx"
    fullpath = os.path.join(pasta, filename)

    doc.save(fullpath)
    return fullpath

def gerar_docx_equipamento_novo(declaracao_id):
    """
    Salva:
    media/declaracoes/ANO-SEMESTRE/PEIA-NOME-equipamento-AAAA-MM.docx
    """
    import os
    from datetime import datetime
    from django.http import Http404
    from django.conf import settings
    from docx import Document
    from decimal import Decimal

    decl = get_object_or_404(declaracao_contrapartida_equipamento, id=declaracao_id)
    itens = decl.itens.all()
    if not itens.exists():
        raise Http404("Declaração de equipamento vazia.")

    ano = decl.ano
    mes = decl.mes
    semestre = 1 if mes <= 6 else 2

    item0 = itens.first()
    projeto_nome = item0.projeto
    projeto_slug = projeto_nome.replace(" ", "_")
    peia = item0.codigo or "SEMPEIA"

    total_geral = sum(Decimal(i.valor_cp or 0) for i in itens)

    path_base = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_equipamento.docx")
    doc = Document(path_base)

    meses = [
        "", "JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO",
        "JULHO","AGOSTO","SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"
    ]

    p = doc.add_paragraph(f"DECLARAÇÃO DE USO DE EQUIPAMENTOS – {projeto_nome}")
    p.runs[0].bold = True

    p = doc.add_paragraph(f"MÊS DE COMPETÊNCIA: {meses[mes]} DE {ano}")
    p.runs[0].bold = True

    doc.add_paragraph("")

    # tabela principal
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "CÓDIGO"
    hdr[1].text = "PROJETO"
    hdr[2].text = "HORAS"
    hdr[3].text = "ATIVIDADES"
    hdr[4].text = "VALOR CP"

    for item in itens:
        row = table.add_row().cells
        row[0].text = item.codigo
        row[1].text = item.projeto
        row[2].text = str(item.horas_alocadas)
        row[3].text = item.descricao or ""
        row[4].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    p = doc.add_paragraph(f"TOTAL GERAL: R$ {total_geral:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))

    # destino
    pasta = os.path.join(settings.MEDIA_ROOT, "declaracoes", f"{ano}-{semestre}")
    os.makedirs(pasta, exist_ok=True)

    filename = f"{peia}-{projeto_slug}-equipamento-{ano}-{mes:02d}.docx"
    fullpath = os.path.join(pasta, filename)

    doc.save(fullpath)
    return fullpath

def _slugify_nome(peia: str | None, nome: str | None) -> str:
    """
    Gera algo como: 'peia1234-projeto_tal'
    """
    peia = (peia or "").strip()
    nome = (nome or "").strip()
    base = f"{peia}-{nome}" if peia else nome
    base = base.lower()
    base = unicodedata.normalize("NFKD", base).encode("ASCII", "ignore").decode("ASCII")
    base = base.replace(" ", "_")
    for ch in ["/", "\\", ":", ";", ","]:
        base = base.replace(ch, "_")
    return base or "projeto_sem_nome"

def _pasta_declaracoes_semestre(ano: int, semestre: int) -> str:
    """
    Ex.: media/declaracoes/2025-1
    """
    pasta = os.path.join(settings.MEDIA_ROOT, "declaracoes", f"{ano}-{semestre}")
    os.makedirs(pasta, exist_ok=True)
    return pasta

def _nome_arquivo_declaracao(ano: int, semestre: int, peia: str, nome: str,
                             tipo: str, mes: int) -> str:
    """
    Ex.: peia1234-projeto_tal-rh-2025-01.docx
    tipo: 'rh', 'pesquisa', 'so'
    """
    slug = _slugify_nome(peia, nome)
    return f"{slug}-{tipo}-{ano}-{mes:02d}.docx"

def _caminho_docx_declaracao(ano: int, semestre: int, peia: str, nome: str,
                             tipo: str, mes: int) -> str:
    pasta = _pasta_declaracoes_semestre(ano, semestre)
    filename = _nome_arquivo_declaracao(ano, semestre, peia, nome, tipo, mes)
    return os.path.join(pasta, filename)

def _caminho_docx_equipamento(ano: int, semestre: int, mes: int) -> str:
    """
    Declaração de equipamento é única por mês:
    equipamentos-AAAA-MM.docx
    """
    pasta = _pasta_declaracoes_semestre(ano, semestre)
    filename = f"equipamentos-{ano}-{mes:02d}.docx"
    return os.path.join(pasta, filename)

try:
    from num2words import num2words
except ImportError:
    raise ImportError("Instale a dependência: pip install num2words")


def valor_por_extenso(valor):
    """
    Recebe float/Decimal e devolve em pt-BR, ex.: 'doze reais e trinta e quatro centavos'
    """
    v = Decimal(valor).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    return num2words(v, to='currency', lang='pt_BR')

def _gerar_docx_rh_para_declaracao(declaracao) -> str:
    """
    Gera o DOCX de RH para uma declaração e salva em disco.
    Retorna o caminho completo do arquivo.
    """
    itens_qs = declaracao_contrapartida_rh_item.objects.filter(declaracao=declaracao)
    if not itens_qs.exists():
        return ""

    primeiro = itens_qs.first()
    ano = primeiro.salario_ano
    mes = primeiro.salario_mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    peia = declaracao.codigo or ""
    total_valor_cp = itens_qs.aggregate(total=Sum('valor_cp'))['total'] or 0

    caminho_docx = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    doc = Document(caminho_docx)

    for p in doc.paragraphs:
        p.text = (
            p.text
            .replace('{{mes_selecionado}}', mes_nome)
            .replace('{{ano_selecionado}}', str(ano))
            .replace('{{nome_projeto}}', projeto_nome)
            .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
        )

    itens = list(itens_qs)

    for paragraph in doc.paragraphs:
        if '{{tabela_itens}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')

            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nome do Bolsista'
            hdr_cells[1].text = 'CPF'
            hdr_cells[2].text = 'Função'
            hdr_cells[3].text = 'Horas alocadas'
            hdr_cells[4].text = 'Salário'
            hdr_cells[5].text = 'Valor CP'
            hdr_cells[6].text = 'Valor Hora'

            for item in itens:
                row_cells = table.add_row().cells
                row_cells[0].text = item.nome
                row_cells[1].text = item.cpf
                row_cells[2].text = item.funcao
                row_cells[3].text = str(item.horas_alocadas)
                row_cells[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row_cells[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row_cells[6].text = f"R$ {item.valor_hora:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            break

    p = doc.add_paragraph('(*) Valor das horas é o produto da multiplicação entre o nº de horas e o quociente da divisão do valor do salário por 160.')
    p = doc.add_paragraph('(**) Mês da competência do contracheque.')
    for _ in range(6):
        doc.add_paragraph('')

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    row = table.rows[0]
    row.cells[0].text = "Anderson Soares"
    row.cells[1].text = "Telma Woerle de Lima Soares"

    row = table.rows[1]
    row.cells[0].text = "Coordenador do Projeto"
    row.cells[1].text = "Diretora da Unidade Embrapii - CEIA/UFG"

    for r in table.rows:
        for cell in r.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'nil')
            tcPr.append(element)

    semestre = 1 if mes <= 6 else 2
    caminho_saida = _caminho_docx_declaracao(ano, semestre, peia, projeto_nome, "rh", mes)
    doc.save(caminho_saida)
    return caminho_saida

def _gerar_docx_so_para_declaracao(declaracao) -> str:
    ano = declaracao.ano
    mes = declaracao.mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    codigo_peia = declaracao.codigo or ''
    total_valor_cp = declaracao.total or 0.0

    caminho_base_preferido = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_so.docx')
    caminho_fallback = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    caminho_docx = caminho_base_preferido if os.path.exists(caminho_base_preferido) else caminho_fallback
    doc = Document(caminho_docx)

    def br_currency(v):
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def apply_replacements_to_text(text):
        return (text
                .replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{codigo_peia}}', codigo_peia)
                .replace('{{valor_total}}', br_currency(total_valor_cp))
                .replace('{{valor_extenso}}', valor_por_extenso(total_valor_cp)))

    for p in doc.paragraphs:
        if p.text:
            p.text = apply_replacements_to_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = apply_replacements_to_text(p.text)

    semestre = 1 if mes <= 6 else 2
    caminho_saida = _caminho_docx_declaracao(ano, semestre, codigo_peia, projeto_nome, "so", mes)
    doc.save(caminho_saida)
    return caminho_saida

def _fmt_moeda_br_decimal(v: Decimal | float | int | None) -> str:
    if v is None:
        v = Decimal("0")
    v = Decimal(v)
    s = f"{v:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {s}"

def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def _set_cell_borders(cell, color="000000", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), size)
        tag.set(qn('w:space'), "0")
        tag.set(qn('w:color'), color)
        tc_borders.append(tag)

def _nota_equipamento_por_nome(nome_equip: str) -> str | None:
    eq = equipamento.objects.filter(nome__iexact=nome_equip).first()
    if not eq:
        return None
    aquis = eq.valor_aquisicao or Decimal("0")
    nos = eq.quantidade_nos or 1
    custo_unid = Decimal(eq.valor_hora or 0)
    unidade = "GPU" if eq.nome.upper().startswith("DGX") else "nó"
    return (
        f"valor de compra original de {_fmt_moeda_br_decimal(aquis)} "
        f"e um custo por {unidade} de {_fmt_moeda_br_decimal(custo_unid)} "
        f"onde o equipamento conta com {nos} {unidade + ('s' if nos != 1 else '')}"
    )

def _gerar_docx_equipamento_para_declaracao(declaracao) -> str:
    itens = declaracao.itens.all().order_by("equipamento", "projeto", "codigo")
    if not itens.exists():
        return ""

    base_path = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_equipamento.docx")
    doc = Document(base_path)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(4.63)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)

    meses = ["", "JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO",
             "JULHO","AGOSTO","SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"]

    p = doc.add_paragraph("DECLARAÇÃO DE USO DE EQUIPAMENTOS")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"MÊS DE COMPETÊNCIA: {meses[declaracao.mes]} DE {declaracao.ano}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")

    col_widths = [Cm(3.5), Cm(5.5), Cm(2.5), Cm(8.0), Cm(3.5)]

    for equip_nome, grupo in groupby(itens, key=lambda i: i.equipamento):
        grupo = list(grupo)
        nota = _nota_equipamento_por_nome(equip_nome)

        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"
        table.autofit = True

        faixa = table.rows[0].cells[0].merge(
            table.rows[0].cells[1]
        ).merge(
            table.rows[0].cells[2]
        ).merge(
            table.rows[0].cells[3]
        ).merge(
            table.rows[0].cells[4]
        )
        _set_cell_shading(faixa, "1F4E79")
        _set_cell_borders(faixa, color="1F4E79")
        par = faixa.paragraphs[0]
        run = par.add_run(f"EQUIPAMENTO: {equip_nome}")
        if nota:
            run.add_text(" *")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

        headers = ["CÓDIGO", "TÍTULO DO PROJETO", "HORAS NO MÊS",
                   "DESCRIÇÃO DAS ATIVIDADES", "VALOR DA CONTRAPARTIDA"]
        for j, text in enumerate(headers):
            c = table.rows[1].cells[j]
            c.width = col_widths[j]
            _set_cell_shading(c, "D9D9D9")
            _set_cell_borders(c)
            par = c.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(text)
            run.bold = True

        total_equip = Decimal("0")
        for item in grupo:
            row = table.add_row()
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
            row.cells[0].text = item.codigo or ""
            row.cells[1].text = item.projeto or ""
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].text = str(item.horas_alocadas or 0)
            row.cells[3].text = item.descricao or ""
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[4].text = _fmt_moeda_br_decimal(item.valor_cp or 0)
            for c in row.cells:
                _set_cell_borders(c)
            total_equip += Decimal(item.valor_cp or 0)

        total_row = table.add_row()
        merged = total_row.cells[0].merge(total_row.cells[1]).merge(
                 total_row.cells[2]).merge(total_row.cells[3])
        _set_cell_shading(merged, "EFEFEF")
        _set_cell_borders(merged)
        merged.paragraphs[0].add_run("TOTAL DA CONTRAPARTIDA").bold = True
        val_cell = total_row.cells[4]
        _set_cell_shading(val_cell, "EFEFEF")
        _set_cell_borders(val_cell)
        par = val_cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        par.add_run(_fmt_moeda_br_decimal(total_equip)).bold = True

        if nota:
            doc.add_paragraph(f"(*) {nota}").runs[0].italic = True

        doc.add_paragraph("")

    total_geral = sum(Decimal(i.valor_cp or 0) for i in itens)
    par = doc.add_paragraph("TOTAL GERAL DA CONTRAPARTIDA: ")
    par.runs[0].bold = True
    par.add_run(_fmt_moeda_br_decimal(total_geral)).bold = True

    semestre = 1 if declaracao.mes <= 6 else 2
    caminho_saida = _caminho_docx_equipamento(declaracao.ano, semestre, declaracao.mes)
    doc.save(caminho_saida)
    return caminho_saida

def _gerar_docx_pesquisa_para_declaracao(declaracao) -> str:
    itens_qs = declaracao.itens.all()
    if not itens_qs.exists():
        return ""

    ano = declaracao.ano
    mes = declaracao.mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    peia = declaracao.codigo or ""

    total_valor_cp = itens_qs.aggregate(total=Sum('valor_cp'))['total'] or 0

    caminho_base_preferido = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_pesquisa.docx')
    caminho_fallback = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    caminho_docx = caminho_base_preferido if os.path.exists(caminho_base_preferido) else caminho_fallback

    doc = Document(caminho_docx)

    for p in doc.paragraphs:
        if p.text:
            p.text = (
                p.text
                .replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
            )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = (
                            p.text
                            .replace('{{mes_selecionado}}', mes_nome)
                            .replace('{{ano_selecionado}}', str(ano))
                            .replace('{{nome_projeto}}', projeto_nome)
                            .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
                        )

    def inserir_tabela_itens(ap_depois):
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Nome do Pesquisador'
        hdr[1].text = 'CPF'
        hdr[2].text = 'Função'
        hdr[3].text = 'Horas alocadas'
        hdr[4].text = 'Salário'
        hdr[5].text = 'Valor CP'

        for item in itens_qs:
            row = table.add_row().cells
            row[0].text = item.nome
            row[1].text = item.cpf
            row[2].text = item.funcao
            row[3].text = str(item.horas_alocadas)
            row[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        ap_depois._element.addnext(table._tbl)

    marcador_encontrado = False
    for paragraph in doc.paragraphs:
        if '{{tabela_itens}}' in paragraph.text:
            marcador_encontrado = True
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')
            inserir_tabela_itens(paragraph)
            break

    if not marcador_encontrado:
        p_anchor = doc.add_paragraph()
        inserir_tabela_itens(p_anchor)

    doc.add_paragraph('(*) Valor das horas é o produto da multiplicação entre o nº de horas e o quociente da divisão do valor do salário por 160.')
    doc.add_paragraph('(**) Mês da competência do contracheque.')

    table_ass = doc.add_table(rows=2, cols=2)
    table_ass.style = 'Table Grid'

    table_ass.rows[0].cells[0].text = "Anderson Soares"
    table_ass.rows[0].cells[1].text = "Telma Woerle de Lima Soares"
    table_ass.rows[1].cells[0].text = "Coordenador do Projeto"
    table_ass.rows[1].cells[1].text = "Diretora da Unidade Embrapii - CEIA/UFG"

    for r in table_ass.rows:
        for cell in r.cells:
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tbl = table_ass._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'nil')
            tcPr.append(element)

    semestre = 1 if mes <= 6 else 2
    caminho_saida = _caminho_docx_declaracao(ano, semestre, peia, projeto_nome, "pesquisa", mes)
    doc.save(caminho_saida)
    return caminho_saida

def gerar_declaracoes_semestre(request):
    """
    Página para:
      - mostrar qtd de projetos no semestre
      - indicar se o semestre está completo (todas as declarações existem)
      - gerar declarações pendentes
      - gerar ZIP com as declarações DOCX existentes
    """
    hoje = date.today()
    ano_atual = hoje.year
    semestre_atual = 1 if hoje.month <= 6 else 2

    # semestre padrão: anterior
    if semestre_atual == 1:
        semestre_default = 2
        ano_default = ano_atual - 1
    else:
        semestre_default = 1
        ano_default = ano_atual

    ano = int(request.GET.get("ano", ano_default))
    semestre = int(request.GET.get("semestre", semestre_default))
    meses_semestre = [1, 2, 3, 4, 5, 6] if semestre == 1 else [7, 8, 9, 10, 11, 12]

    # projetos ativos no semestre
    data_ini = date(ano, 1 if semestre == 1 else 7, 1)
    mes_fim = 6 if semestre == 1 else 12
    dia_fim = 30 if semestre == 1 else 31
    data_fim = date(ano, mes_fim, dia_fim)

    projetos = projeto.objects.filter(
        ativo=True,
        data_inicio__lte=data_fim,
        data_fim__gte=data_ini
    ).order_by("nome")

    projetos_qtd = projetos.count()

    # ---- contagem de declarações existentes ----
    total_existente = 0
    total_esperado = (projetos_qtd * 3 * len(meses_semestre)) + len(meses_semestre)

    for p in projetos:
        for mes in meses_semestre:
            if declaracao_contrapartida_pesquisa.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                total_existente += 1
            if declaracao_contrapartida_rh.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                total_existente += 1
            if declaracao_contrapartida_so.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                total_existente += 1

    for mes in meses_semestre:
        if declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).exists():
            total_existente += 1

    semestre_completo = (total_existente == total_esperado)

    # ---- gerar declarações pendentes ----
    if request.GET.get("gerar") == "1":
        total_pesq = total_rh = total_so = total_equip = 0

        for p in projetos:
            for mes in meses_semestre:

                # PESQUISA
                if not declaracao_contrapartida_pesquisa.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                    gerar_declaracao_contrapartida_pesquisa(request, p.id, mes, ano)
                decl_p = declaracao_contrapartida_pesquisa.objects.filter(id_projeto=p.id, mes=mes, ano=ano).first()
                if decl_p:
                    _gerar_docx_pesquisa_para_declaracao(decl_p)
                    total_pesq += 1

                # RH
                if not declaracao_contrapartida_rh.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                    gerar_declaracao_contrapartida_rh(request, p.id, mes, ano)
                decl_rh = declaracao_contrapartida_rh.objects.filter(id_projeto=p.id, mes=mes, ano=ano).first()
                if decl_rh:
                    _gerar_docx_rh_para_declaracao(decl_rh)
                    total_rh += 1

                # SO
                if not declaracao_contrapartida_so.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                    gerar_declaracao_contrapartida_so(request, p.id, mes, ano)
                decl_so = declaracao_contrapartida_so.objects.filter(id_projeto=p.id, mes=mes, ano=ano).first()
                if decl_so:
                    _gerar_docx_so_para_declaracao(decl_so)
                    total_so += 1

        # EQUIPAMENTO (1 por mês, global)
        for mes in meses_semestre:
            if not declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).exists():
                gerar_declaracao_contrapartida_equipamento(request, p.id, mes, ano)
            decl_eq = declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).first()
            if decl_eq:
                _gerar_docx_equipamento_para_declaracao(decl_eq)
                total_equip += 1

        messages.success(
            request,
            f"Declarações (e DOCX) gerados. Pesquisa={total_pesq}, RH={total_rh}, SO={total_so}, Equip={total_equip}"
        )

        request.session["ultima_geracao_zip"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        return redirect(reverse("gerar_declaracoes_semestre") + f"?ano={ano}&semestre={semestre}")

    # ---- ZIP do semestre ----
    pasta_semestre = _pasta_declaracoes_semestre(ano, semestre)
    zip_filename = f"declaracoes_{ano}-{semestre}.zip"
    zip_path = os.path.join(pasta_semestre, zip_filename)
    zip_exists = os.path.exists(zip_path)

    if request.GET.get("zip") == "1":
        # recria o ZIP sempre, com o que existir
        if os.path.exists(zip_path):
            os.remove(zip_path)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            # 1) arquivos por projeto
            for p in projetos:
                peia = p.peia
                nome_proj = p.nome

                for mes in meses_semestre:
                    caminho_rh = _caminho_docx_declaracao(ano, semestre, peia, nome_proj, "rh", mes)
                    if os.path.exists(caminho_rh):
                        zipf.write(caminho_rh, os.path.relpath(caminho_rh, pasta_semestre))

                    caminho_pesq = _caminho_docx_declaracao(ano, semestre, peia, nome_proj, "pesquisa", mes)
                    if os.path.exists(caminho_pesq):
                        zipf.write(caminho_pesq, os.path.relpath(caminho_pesq, pasta_semestre))

                    caminho_so = _caminho_docx_declaracao(ano, semestre, peia, nome_proj, "so", mes)
                    if os.path.exists(caminho_so):
                        zipf.write(caminho_so, os.path.relpath(caminho_so, pasta_semestre))

            # 2) equipamentos (1 por mês, global) — fora do loop de projetos
            for mes in meses_semestre:
                caminho_eq = _caminho_docx_equipamento(ano, semestre, mes)
                if os.path.exists(caminho_eq):
                    zipf.write(caminho_eq, os.path.relpath(caminho_eq, pasta_semestre))



        messages.success(request, "ZIP criado com sucesso!")
        return redirect(reverse("gerar_declaracoes_semestre") + f"?ano={ano}&semestre={semestre}")

    ultima_data = request.session.get("ultima_geracao_zip")

    return render(request, "declaracao/gerar_declaracoes_semestre.html", {
        "ano": ano,
        "semestre": semestre,
        "projetos_qtd": projetos_qtd,
        "semestre_completo": semestre_completo,
        "zip_exists": os.path.exists(zip_path),
        "zip_filename": zip_filename,
        "ultima_data": ultima_data,
    })


# =========================================================
# PROGRESSO GLOBAL (DEIXE SÓ ESTE, UMA ÚNICA VEZ NO ARQUIVO)
# =========================================================
progresso_lock = Lock()
progresso = {
    "status": "aguardando",   # aguardando | gerando | finalizado | erro
    "percentual": 0,
    "mensagem": "",
    "ano": None,
    "semestre": None,
    "finalizado": False,
}

def _set_progresso(**kwargs):
    with progresso_lock:
        progresso.update(kwargs)

def _no_cache_json(data, status=200):
    resp = JsonResponse(data, status=status)
    resp["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp["Pragma"] = "no-cache"
    resp["Expires"] = "0"
    return resp

@never_cache
@login_required
def progresso_semestre(request):
    with progresso_lock:
        return _no_cache_json(progresso)

@never_cache
@login_required
def gerar_semestre_ajax(request):
    """
    Gera declarações + DOCX no disco, atualizando o dict global 'progresso'.
    O front faz polling em /progresso_semestre/ pra atualizar a barra.
    """
    try:
        ano = int(request.GET.get("ano"))
        semestre = int(request.GET.get("semestre"))
    except (TypeError, ValueError):
        _set_progresso(status="erro", mensagem="Parâmetros inválidos (ano/semestre).", finalizado=True)
        return _no_cache_json({"ok": False, "erro": "Parâmetros inválidos."}, status=400)

    meses = [1, 2, 3, 4, 5, 6] if semestre == 1 else [7, 8, 9, 10, 11, 12]

    data_ini = date(ano, 1 if semestre == 1 else 7, 1)
    data_fim = date(ano, 6 if semestre == 1 else 12, 30 if semestre == 1 else 31)

    projetos = projeto.objects.filter(
        ativo=True,
        data_inicio__lte=data_fim,
        data_fim__gte=data_ini
    ).order_by("nome")

    projetos_qtd = projetos.count()

    # total de “passos” pra barra:
    # 3 docx por projeto/mês (pesq, rh, so) + 1 docx de equipamento por mês
    total_ops = (projetos_qtd * 3 * len(meses)) + len(meses)
    if total_ops <= 0:
        _set_progresso(
            status="finalizado",
            percentual=100,
            mensagem="Nada a gerar (sem projetos no semestre).",
            ano=ano,
            semestre=semestre,
            finalizado=True,
        )
        return _no_cache_json({"ok": True, "vazio": True})

    atual = 0

    _set_progresso(
        status="gerando",
        percentual=0,
        mensagem="Iniciando geração...",
        ano=ano,
        semestre=semestre,
        finalizado=False,
    )
    def _msg_projeto(tipo: str, ano: int, mes: int, p) -> str:
        return f"Gerando {tipo} - {ano} - {mes:02d} - {p.peia} - {p.nome}"

    def _msg_equip(tipo: str, ano: int, mes: int) -> str:
        return f"Gerando {tipo} - {ano} - {mes:02d}"


    # log opcional (ajuda a debugar travas)
    logdir = os.path.join(settings.MEDIA_ROOT, "logs")
    os.makedirs(logdir, exist_ok=True)
    logpath = os.path.join(logdir, f"geracao_{ano}_{semestre}.log")

    def log(msg):
        with open(logpath, "a", encoding="utf-8") as f:
            f.write(f"{datetime.now()} - {msg}\n")

    log(f"Início geração semestre {semestre}/{ano}, projetos={projetos_qtd}")

    # pega um projeto “fallback” para a declaração de equipamento (sua função exige projeto_id)
    p_fallback = projetos.first()

    try:
        # -------------------------
        # LOOP PROJETOS / MESES
        # -------------------------
        for p in projetos:
            for mes in meses:

                # ===== PESQUISA =====
                if not declaracao_contrapartida_pesquisa.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                    gerar_declaracao_contrapartida_pesquisa(request, p.id, mes, ano)

                decl_p = declaracao_contrapartida_pesquisa.objects.filter(id_projeto=p.id, mes=mes, ano=ano).first()
                if decl_p:
                    _gerar_docx_pesquisa_para_declaracao(decl_p)
                    log(f"Pesquisa: {p.nome} {mes}/{ano}")

                atual += 1
                _set_progresso(
                    percentual=int(100 * atual / total_ops),
                    mensagem=_msg_projeto("Pesquisa", ano, mes, p)
                )


                # ===== RH =====
                if not declaracao_contrapartida_rh.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                    gerar_declaracao_contrapartida_rh(request, p.id, mes, ano)

                decl_rh = declaracao_contrapartida_rh.objects.filter(id_projeto=p.id, mes=mes, ano=ano).first()
                if decl_rh:
                    _gerar_docx_rh_para_declaracao(decl_rh)
                    log(f"RH: {p.nome} {mes}/{ano}")

                atual += 1
                _set_progresso(
                    percentual=int(100 * atual / total_ops),
                    mensagem=_msg_projeto("RH", ano, mes, p)
                )


                # ===== SO =====
                if not declaracao_contrapartida_so.objects.filter(id_projeto=p.id, mes=mes, ano=ano).exists():
                    gerar_declaracao_contrapartida_so(request, p.id, mes, ano)

                decl_so = declaracao_contrapartida_so.objects.filter(id_projeto=p.id, mes=mes, ano=ano).first()
                if decl_so:
                    _gerar_docx_so_para_declaracao(decl_so)
                    log(f"SO: {p.nome} {mes}/{ano}")

                atual += 1
                _set_progresso(
                    percentual=int(100 * atual / total_ops),
                    mensagem=_msg_projeto("SO", ano, mes, p)
                )


        # -------------------------
        # EQUIPAMENTOS (1 por mês)
        # -------------------------
        for mes in meses:
            if not declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).exists():
                if p_fallback:
                    gerar_declaracao_contrapartida_equipamento(request, p_fallback.id, mes, ano)

            decl_eq = declaracao_contrapartida_equipamento.objects.filter(mes=mes, ano=ano).first()
            if decl_eq:
                _gerar_docx_equipamento_para_declaracao(decl_eq)
                log(f"Equipamento: {mes}/{ano}")

            atual += 1
            _set_progresso(
                percentual=int(100 * atual / total_ops),
                mensagem=_msg_equip("Equipamento", ano, mes)
            )


        _set_progresso(status="finalizado", percentual=100, mensagem="Geração concluída!", finalizado=True)
        log("Finalizado com sucesso.")
        return _no_cache_json({"ok": True})

    except Exception as e:
        log(f"ERRO: {repr(e)}")
        _set_progresso(status="erro", mensagem=f"Erro: {e}", finalizado=True)
        return _no_cache_json({"ok": False, "erro": str(e)}, status=500)
