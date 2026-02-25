import os
import unicodedata
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from itertools import groupby

from django.conf import settings
from django.db.models import Sum

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

from .models import (
    declaracao_contrapartida_rh,
    declaracao_contrapartida_rh_item,
    declaracao_contrapartida_pesquisa,
    declaracao_contrapartida_so,
    declaracao_contrapartida_equipamento,
    declaracao_contrapartida_equipamento_item,
)
from contrapartida.models import equipamento

try:
    from num2words import num2words
except ImportError:
    raise ImportError("Instale a dependência: pip install num2words")


# =========================================================
# HELPERS
# =========================================================

def valor_por_extenso(valor):
    """
    Recebe float/Decimal e devolve em pt-BR, ex.: 'doze reais e trinta e quatro centavos'
    """
    v = Decimal(valor).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    return num2words(v, to='currency', lang='pt_BR')


def _fmt_moeda_br_decimal(v) -> str:
    if v is None:
        v = Decimal("0")
    v = Decimal(v)
    s = f"{v:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {s}"


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color="000000", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), size)
        tag.set(qn('w:space'), "0")
        tag.set(qn('w:color'), color)
        tc_borders.append(tag)


def _nota_equipamento_por_nome(nome_equip: str):
    eq = equipamento.objects.filter(nome__iexact=nome_equip).first()
    if not eq:
        return None
    aquis = eq.valor_aquisicao or Decimal("0")
    nos = eq.quantidade_nos or 1
    custo_unid = Decimal(eq.valor_hora or 0)
    unidade = "GPU" if eq.nome.upper().startswith("DGX") else "nó"
    return (
        f"valor de compra original de {_fmt_moeda_br_decimal(aquis)} "
        f"e um custo por {unidade} de {_fmt_moeda_br_decimal(custo_unid)} "
        f"onde o equipamento conta com {nos} {unidade + ('s' if nos != 1 else '')}"
    )


# =========================================================
# PATH HELPERS
# =========================================================

def _slugify_nome(peia, nome) -> str:
    """
    Gera algo como: 'peia1234-projeto_tal'
    """
    peia = (peia or "").strip()
    nome = (nome or "").strip()
    base = f"{peia}-{nome}" if peia else nome
    base = base.lower()
    base = unicodedata.normalize("NFKD", base).encode("ASCII", "ignore").decode("ASCII")
    base = base.replace(" ", "_")
    for ch in ["/", "\\", ":", ";", ","]:
        base = base.replace(ch, "_")
    return base or "projeto_sem_nome"


def _pasta_declaracoes_semestre(ano: int, semestre: int) -> str:
    """
    Ex.: media/declaracoes/2025-1
    """
    pasta = os.path.join(settings.MEDIA_ROOT, "declaracoes", f"{ano}-{semestre}")
    os.makedirs(pasta, exist_ok=True)
    return pasta


def _nome_arquivo_declaracao(ano: int, semestre: int, peia: str, nome: str,
                             tipo: str, mes: int) -> str:
    """
    Ex.: peia1234-projeto_tal-rh-2025-01.docx
    tipo: 'rh', 'pesquisa', 'so'
    """
    slug = _slugify_nome(peia, nome)
    return f"{slug}-{tipo}-{ano}-{mes:02d}.docx"


def _caminho_docx_declaracao(ano: int, semestre: int, peia: str, nome: str,
                             tipo: str, mes: int) -> str:
    pasta = _pasta_declaracoes_semestre(ano, semestre)
    filename = _nome_arquivo_declaracao(ano, semestre, peia, nome, tipo, mes)
    return os.path.join(pasta, filename)


def _caminho_docx_equipamento(ano: int, semestre: int, mes: int) -> str:
    """
    Declaração de equipamento é única por mês:
    equipamentos-AAAA-MM.docx
    """
    pasta = _pasta_declaracoes_semestre(ano, semestre)
    filename = f"equipamentos-{ano}-{mes:02d}.docx"
    return os.path.join(pasta, filename)


# =========================================================
# RH
# =========================================================

def gerar_docx_rh(declaracao) -> str:
    """
    Gera o DOCX de RH para uma declaração e salva em disco.
    Retorna o caminho completo do arquivo.
    """
    itens_qs = declaracao_contrapartida_rh_item.objects.filter(declaracao=declaracao)
    if not itens_qs.exists():
        return ""

    primeiro = itens_qs.first()
    ano = primeiro.salario_ano
    mes = primeiro.salario_mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    peia = declaracao.codigo or ""
    total_valor_cp = itens_qs.aggregate(total=Sum('valor_cp'))['total'] or 0

    caminho_docx = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    doc = Document(caminho_docx)

    for p in doc.paragraphs:
        p.text = (
            p.text
            .replace('{{mes_selecionado}}', mes_nome)
            .replace('{{ano_selecionado}}', str(ano))
            .replace('{{nome_projeto}}', projeto_nome)
            .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
        )

    itens = list(itens_qs)

    for paragraph in doc.paragraphs:
        if '{{tabela_itens}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')

            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nome do Bolsista'
            hdr_cells[1].text = 'CPF'
            hdr_cells[2].text = 'Função'
            hdr_cells[3].text = 'Horas alocadas'
            hdr_cells[4].text = 'Salário'
            hdr_cells[5].text = 'Valor CP'
            hdr_cells[6].text = 'Valor Hora'

            for item in itens:
                row_cells = table.add_row().cells
                row_cells[0].text = item.nome
                row_cells[1].text = item.cpf
                row_cells[2].text = item.funcao
                row_cells[3].text = str(item.horas_alocadas)
                row_cells[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row_cells[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                row_cells[6].text = f"R$ {item.valor_hora:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            break

    doc.add_paragraph('(*) Valor das horas é o produto da multiplicação entre o nº de horas e o quociente da divisão do valor do salário por 160.')
    doc.add_paragraph('(**) Mês da competência do contracheque.')
    for _ in range(6):
        doc.add_paragraph('')

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    row = table.rows[0]
    row.cells[0].text = "Anderson Soares"
    row.cells[1].text = "Telma Woerle de Lima Soares"

    row = table.rows[1]
    row.cells[0].text = "Coordenador do Projeto"
    row.cells[1].text = "Diretora da Unidade Embrapii - CEIA/UFG"

    for r in table.rows:
        for cell in r.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'nil')
            tcPr.append(element)

    semestre = 1 if mes <= 6 else 2
    caminho_saida = _caminho_docx_declaracao(ano, semestre, peia, projeto_nome, "rh", mes)
    doc.save(caminho_saida)
    return caminho_saida


# =========================================================
# SO
# =========================================================

def gerar_docx_so(declaracao) -> str:
    ano = declaracao.ano
    mes = declaracao.mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    codigo_peia = declaracao.codigo or ''
    total_valor_cp = declaracao.total or 0.0

    caminho_base_preferido = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_so.docx')
    caminho_fallback = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    caminho_docx = caminho_base_preferido if os.path.exists(caminho_base_preferido) else caminho_fallback
    doc = Document(caminho_docx)

    def br_currency(v):
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def apply_replacements_to_text(text):
        return (text
                .replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{codigo_peia}}', codigo_peia)
                .replace('{{valor_total}}', br_currency(total_valor_cp))
                .replace('{{valor_extenso}}', valor_por_extenso(total_valor_cp)))

    for p in doc.paragraphs:
        if p.text:
            p.text = apply_replacements_to_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = apply_replacements_to_text(p.text)

    semestre = 1 if mes <= 6 else 2
    caminho_saida = _caminho_docx_declaracao(ano, semestre, codigo_peia, projeto_nome, "so", mes)
    doc.save(caminho_saida)
    return caminho_saida


# =========================================================
# PESQUISA
# =========================================================

def gerar_docx_pesquisa(declaracao) -> str:
    itens_qs = declaracao.itens.all()
    if not itens_qs.exists():
        return ""

    ano = declaracao.ano
    mes = declaracao.mes
    mes_nome = datetime(ano, mes, 1).strftime('%B').capitalize()
    projeto_nome = declaracao.projeto
    peia = declaracao.codigo or ""

    total_valor_cp = itens_qs.aggregate(total=Sum('valor_cp'))['total'] or 0

    caminho_base_preferido = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_pesquisa.docx')
    caminho_fallback = os.path.join(settings.BASE_DIR, 'contrapartida', 'static', 'base_rh.docx')
    caminho_docx = caminho_base_preferido if os.path.exists(caminho_base_preferido) else caminho_fallback

    doc = Document(caminho_docx)

    for p in doc.paragraphs:
        if p.text:
            p.text = (
                p.text
                .replace('{{mes_selecionado}}', mes_nome)
                .replace('{{ano_selecionado}}', str(ano))
                .replace('{{nome_projeto}}', projeto_nome)
                .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
            )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = (
                            p.text
                            .replace('{{mes_selecionado}}', mes_nome)
                            .replace('{{ano_selecionado}}', str(ano))
                            .replace('{{nome_projeto}}', projeto_nome)
                            .replace('{{valor_total}}', f"R$ {total_valor_cp:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."))
                        )

    def inserir_tabela_itens(ap_depois):
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Nome do Pesquisador'
        hdr[1].text = 'CPF'
        hdr[2].text = 'Função'
        hdr[3].text = 'Horas alocadas'
        hdr[4].text = 'Salário'
        hdr[5].text = 'Valor CP'

        for item in itens_qs:
            row = table.add_row().cells
            row[0].text = item.nome
            row[1].text = item.cpf
            row[2].text = item.funcao
            row[3].text = str(item.horas_alocadas)
            row[4].text = f"R$ {item.salario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row[5].text = f"R$ {item.valor_cp:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        ap_depois._element.addnext(table._tbl)

    marcador_encontrado = False
    for paragraph in doc.paragraphs:
        if '{{tabela_itens}}' in paragraph.text:
            marcador_encontrado = True
            paragraph.text = paragraph.text.replace('{{tabela_itens}}', '')
            inserir_tabela_itens(paragraph)
            break

    if not marcador_encontrado:
        p_anchor = doc.add_paragraph()
        inserir_tabela_itens(p_anchor)

    doc.add_paragraph('(*) Valor das horas é o produto da multiplicação entre o nº de horas e o quociente da divisão do valor do salário por 160.')
    doc.add_paragraph('(**) Mês da competência do contracheque.')

    table_ass = doc.add_table(rows=2, cols=2)
    table_ass.style = 'Table Grid'

    table_ass.rows[0].cells[0].text = "Anderson Soares"
    table_ass.rows[0].cells[1].text = "Telma Woerle de Lima Soares"
    table_ass.rows[1].cells[0].text = "Coordenador do Projeto"
    table_ass.rows[1].cells[1].text = "Diretora da Unidade Embrapii - CEIA/UFG"

    for r in table_ass.rows:
        for cell in r.cells:
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tbl = table_ass._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'nil')
            tcPr.append(element)

    semestre = 1 if mes <= 6 else 2
    caminho_saida = _caminho_docx_declaracao(ano, semestre, peia, projeto_nome, "pesquisa", mes)
    doc.save(caminho_saida)
    return caminho_saida


# =========================================================
# EQUIPAMENTO
# =========================================================

def gerar_docx_equipamento(declaracao) -> str:
    itens = declaracao.itens.all().order_by("equipamento", "projeto", "codigo")
    if not itens.exists():
        return ""

    base_path = os.path.join(settings.BASE_DIR, "contrapartida", "static", "base_equipamento.docx")
    doc = Document(base_path)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(4.63)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)

    meses = ["", "JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO",
             "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"]

    p = doc.add_paragraph("DECLARAÇÃO DE USO DE EQUIPAMENTOS")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"MÊS DE COMPETÊNCIA: {meses[declaracao.mes]} DE {declaracao.ano}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")

    col_widths = [Cm(3.5), Cm(5.5), Cm(2.5), Cm(8.0), Cm(3.5)]

    for equip_nome, grupo in groupby(itens, key=lambda i: i.equipamento):
        grupo = list(grupo)
        nota = _nota_equipamento_por_nome(equip_nome)

        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"
        table.autofit = True

        faixa = table.rows[0].cells[0].merge(
            table.rows[0].cells[1]
        ).merge(
            table.rows[0].cells[2]
        ).merge(
            table.rows[0].cells[3]
        ).merge(
            table.rows[0].cells[4]
        )
        _set_cell_shading(faixa, "1F4E79")
        _set_cell_borders(faixa, color="1F4E79")
        par = faixa.paragraphs[0]
        run = par.add_run(f"EQUIPAMENTO: {equip_nome}")
        if nota:
            run.add_text(" *")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

        headers = ["CÓDIGO", "TÍTULO DO PROJETO", "HORAS NO MÊS",
                   "DESCRIÇÃO DAS ATIVIDADES", "VALOR DA CONTRAPARTIDA"]
        for j, text in enumerate(headers):
            c = table.rows[1].cells[j]
            c.width = col_widths[j]
            _set_cell_shading(c, "D9D9D9")
            _set_cell_borders(c)
            par = c.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(text)
            run.bold = True

        total_equip = Decimal("0")
        for item in grupo:
            row = table.add_row()
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
            row.cells[0].text = item.codigo or ""
            row.cells[1].text = item.projeto or ""
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].text = str(item.horas_alocadas or 0)
            row.cells[3].text = item.descricao or ""
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[4].text = _fmt_moeda_br_decimal(item.valor_cp or 0)
            for c in row.cells:
                _set_cell_borders(c)
            total_equip += Decimal(item.valor_cp or 0)

        total_row = table.add_row()
        merged = total_row.cells[0].merge(total_row.cells[1]).merge(
                 total_row.cells[2]).merge(total_row.cells[3])
        _set_cell_shading(merged, "EFEFEF")
        _set_cell_borders(merged)
        merged.paragraphs[0].add_run("TOTAL DA CONTRAPARTIDA").bold = True
        val_cell = total_row.cells[4]
        _set_cell_shading(val_cell, "EFEFEF")
        _set_cell_borders(val_cell)
        par = val_cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        par.add_run(_fmt_moeda_br_decimal(total_equip)).bold = True

        if nota:
            doc.add_paragraph(f"(*) {nota}").runs[0].italic = True

        doc.add_paragraph("")

    total_geral = sum(Decimal(i.valor_cp or 0) for i in itens)
    par = doc.add_paragraph("TOTAL GERAL DA CONTRAPARTIDA: ")
    par.runs[0].bold = True
    par.add_run(_fmt_moeda_br_decimal(total_geral)).bold = True

    semestre = 1 if declaracao.mes <= 6 else 2
    caminho_saida = _caminho_docx_equipamento(declaracao.ano, semestre, declaracao.mes)
    doc.save(caminho_saida)
    return caminho_saida
